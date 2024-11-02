import os
from flask_wtf import CSRFProtect, FlaskForm
from datetime import datetime, timezone
import humanize
from sqlalchemy.exc import IntegrityError
from wtforms.validators import DataRequired, Email, EqualTo,Length,ValidationError,Optional,Regexp
from flask_wtf.csrf import generate_csrf, validate_csrf,CSRFError
from wtforms import SubmitField
from wtforms import StringField, SubmitField, FileField, IntegerField,FloatField,PasswordField,SelectField,TextAreaField,BooleanField,  DateField,FileField
from wtforms.validators import DataRequired
from flask import Flask, send_file,render_template, request, redirect, url_for, flash,jsonify, abort,session
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, login_user, login_required, logout_user, UserMixin,login_manager,current_user
from Crypto.Hash import SHA256
from flask_migrate import Migrate
from werkzeug.datastructures import MultiDict,FileStorage
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
from alembic import op
import re
from functools import wraps
from flask import session,abort
from sqlalchemy import func,String
# from sqlalchemy.dialects.postgresql import UUID
import uuid
import hashlib
from uuid import UUID
from sqlalchemy import Column, UUID
import json
from sqlalchemy import ForeignKey
import pandas as pd
import pymysql, logging
from dotenv import load_dotenv
from flask_wtf.file import FileAllowed, FileRequired, FileField,FileSize
# from utils import load_config, generate_db_uri
import random, time
from datetime import timedelta
from dateutil import parser
from docx import Document
from io import BytesIO
from flask_session import Session

# Load environment variables
load_dotenv()

# Configuration
DEFAULT_ADMIN_EMAIL = os.getenv('DEFAULT_ADMIN_EMAIL')
DEFAULT_ADMIN_PASSWORD = os.getenv('DEFAULT_ADMIN_PASSWORD')


app = Flask(__name__)
#incase of deployment to live server
application = app

#sqlite3 flask default db
app.config['SECRET_KEY'] = 'your_secret_key_here'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users.db' 
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20 MB
app.config['WTF_CSRF_ENABLED']= True


app.config['SESSION_TYPE'] = 'filesystem'  # Or 'sqlalchemy' for database session
app.config['SESSION_FILE_DIR'] = './flask_sessions/'  # Ensure this path exists
app.config['SESSION_PERMANENT'] = False  # Set to false if you don't want sessions to persist forever


Session(app)  # Initialize server-side sessions

csrf = CSRFProtect(app)
db = SQLAlchemy(app)
migrate = Migrate(app, db)

login_manager = LoginManager()
login_manager.init_app(app)
 
@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(error):
    return "File is too large. The maximum file size is 20MB.", 413


# Create the default admin
def create_default_admin():
    default_admin = AppAdmin.query.filter_by(email=DEFAULT_ADMIN_EMAIL).first()
    if not default_admin:
        hashed_password = generate_password_hash(DEFAULT_ADMIN_PASSWORD, method='pbkdf2:sha256')
         # Fetch the last serial number
        last_serial = db.session.query(func.max(AppAdmin.serial_number)).scalar() or 0

        new_admin = AppAdmin(
            email=DEFAULT_ADMIN_EMAIL,
            password=hashed_password,
            
            # Assign the next available serial number
            serial_number=last_serial + 1,
            admin_name='Default Super Admin'
            
        )
        db.session.add(new_admin)
        db.session.commit()
        print("Default super admin created")
    else:
        print("Default super admin already exists")


##############################################
session_dir = './flask_sessions/'

for file in os.listdir(session_dir):
    file_path = os.path.join(session_dir, file)
    os.remove(file_path)

#######################################


@login_manager.user_loader
def load_user(user_id):
    user_type, id_str = user_id.split("_")
    
    if user_type == "farmer":
        return Farmer.query.get(UUID(id_str))  # Convert to UUID if necessary
    elif user_type == "admin":
        return Admin.query.get(UUID(id_str))
    elif user_type == "app_admin":
        return AppAdmin.query.get(UUID(id_str))
    
    return None

#############################################################

# def role_required(*roles):
#     def decorator(f):
#         @wraps(f)
#         def decorated_function(*args, **kwargs):
#             if not (current_user.is_authenticated and any(getattr(current_user, role, False) for role in roles)):
#                 flash('You do not have access to this resource.', 'danger')
#                 return redirect(url_for('login'))
#             return f(*args, **kwargs)
#         return decorated_function
#     return decorator

def app_admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Check if the current user is an Admin or AppAdmin
        if not isinstance(current_user, AppAdmin):
            abort(403)  # Forbidden access
        return f(*args, **kwargs)
    return decorated_function


def admin_or_app_admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Check if the current user is an Admin or AppAdmin
        if not (isinstance(current_user, Admin) or isinstance(current_user, AppAdmin)):
            abort(403)  # Forbidden access
        return f(*args, **kwargs)
    return decorated_function


def all_users_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Check if the user is authenticated
        if not current_user.is_authenticated:
            abort(403)  # Forbidden access
        return f(*args, **kwargs)
    return decorated_function

############################################################
#...............flask_form...........................................
class LoginForm(FlaskForm):
    email_or_phone = StringField('Email or Phone Number', validators=[DataRequired()])

    def validate_email(self, field):
        value = field.data
        # Regex for email
        email_regex = r"[^@]+@[^@]+\.[^@]+"
        # Regex for phone number (optional: adjust based on the format you expect)
        phone_regex = r"^\+?\d{10,15}$"  # allows international numbers like +234 or 09078780989

        if not (re.match(email_regex, value) or re.match(phone_regex, value)):
            raise ValidationError('Please enter a valid email address or phone number.')

    password = PasswordField('Password', validators=[DataRequired()])
    remember_me = BooleanField('Remember Me')
    submit = SubmitField('Login')

    @classmethod
    def from_json(cls, data):
        # Implement logic to create a LoginForm object from JSON data
        email_or_phone = data.get('email_or_phone')
        password = data.get('password')
        return cls(email_or_phone=email_or_phone, password=password)
    

#...........flask_search_form................#

class SearchAdminForm(FlaskForm):
    admin_search = StringField('Search', validators=[DataRequired()])
    submit = SubmitField('Search')


class SearchFarmerForm(FlaskForm):
    farmer_search = StringField('Search', validators=[DataRequired()])
    submit = SubmitField('Search')


class FarmerForm(FlaskForm):
    # Personal Information
    farmer_firstname = StringField('First Name', validators=[DataRequired(), Length(min=2, max=150)])
    farmer_middlename = StringField('Middle Name', validators=[Optional(), Length(min=2, max=150)])
    farmer_lastname = StringField('Last Name', validators=[DataRequired(), Length(min=2, max=150)])
    password = PasswordField('Password', validators=[Optional()])
    farmer_gender = SelectField('Gender', choices=[('M', 'Male'), ('F', 'Female')], validators=[DataRequired()])
    
    # Additional Information
    date_of_birth = DateField('Date of Birth', validators=[DataRequired()])
    phone_number = StringField('Phone Number', validators=[
        DataRequired(),
        Regexp(r'^(\+234|0)\d{10}$', message="Phone number must start with '+234' or '0' followed by 10 digits.")
    ])
    email = StringField('Email', validators=[DataRequired(), Email()])
    bvn_number = StringField('BVN Number', validators=[DataRequired(), Length(min=11, max=11)])
    id_card_number = StringField('ID Card Number', validators=[DataRequired()])
    id_card_type = SelectField('ID Card Type', choices=[('National ID', 'National ID'), ('Voter\'s Card', 'Voter\'s Card'), ('Driver\'s License', 'Driver\'s License'), ('International Passport', 'International Passport')])
    farm_size = StringField('Farm Size (HA)', validators=[DataRequired()])
    farm_location = StringField('Farm Location', validators=[DataRequired()])  # Add this line
    gps_coordinates = StringField('GPS Coordinates')
    ward = StringField('Ward', validators=[DataRequired()])
    local_govt_council = StringField('Local Govt Council', validators=[DataRequired()])
    # Image Upload
    passport_photo = FileField('Passport Photograph', 
    validators=[
        Optional(), 
        FileAllowed(['jpg', 'png', 'gif'], 'Only .jpg, .png, and .gif formats are allowed'), 
        FileSize(max_size=0.5 * 1024 * 1024, message='File size exceeds 500kb')
    ]
)
    submit = SubmitField('Submit')
    
    def validate_email(self, email):
        farmer = Farmer.query.filter_by(email=email.data).first()
        if farmer:
            raise ValidationError('This email is already registered.')

    def validate_bvn_number(self, bvn_number):
        farmer = Farmer.query.filter_by(bvn_number=bvn_number.data).first()
        if farmer:
            raise ValidationError('This BVN number is already registered.')

    def validate_id_card_number(self, id_card_number):
        farmer = Farmer.query.filter_by(id_card_number=id_card_number.data).first()
        if farmer:
            raise ValidationError('This ID card number is already registered.')
        

# Custom validator to further enforce correct length
    def validate_phone_number(form, field):
        phone_number = field.data
        if phone_number.startswith('+234'):
            if len(phone_number) != 14:
                raise ValidationError("Phone number must be 13 digits after +234.")
        elif phone_number.startswith('0'):
            if len(phone_number) != 11:
                raise ValidationError("Phone number must be 11 digits including the leading 0.")
 
class EditFarmerForm(FlaskForm):
   
    farmer_firstname = StringField('First Name', validators=[DataRequired()])
    farmer_lastname = StringField('Last Name', validators=[DataRequired()])
    email = StringField('Email', validators=[DataRequired(), Email()])
    phone_number = StringField('Phone Number', validators=[
        DataRequired(),
        Regexp(r'^(\+234|0)\d{10}$', message="Phone number must start with '+234' or '0' followed by 10 digits.")
    ])
    farm_location = StringField('Farm Location', validators=[DataRequired()])
    farm_size = StringField('Farm Size (hectares)', validators=[DataRequired()])
    gps_coordinates = StringField('GPS Coordinates')
    id_card_type = SelectField('ID Card Type', choices=[('National ID', 'National ID'), ('Voter\'s Card', 'Voter\'s Card'), ('Driver\'s License', 'Driver\'s License'), ('International Passport', 'International Passport')])
    id_card_number = StringField('ID Card Number')
    bvn_number = StringField('BVN Number', validators=[ DataRequired(),Regexp(r'^\d{11}$', message="BVN must be exactly 11 digits.")
    ])
    passport_photo = FileField('Passport Photograph', 
    validators=[
        Optional(), 
        FileAllowed(['jpg', 'png', 'gif'], 'Only .jpg, .png, and .gif formats are allowed'), 
        FileSize(max_size=0.5 * 1024 * 1024, message='File size exceeds 500kb')
    ])
    submit = SubmitField('Update Details')


class AdminForm(FlaskForm):

    admin_name = StringField('Admin Name', validators=[DataRequired(), Length(min=2, max=50)])
    phone_number = StringField('Phone Number', validators=[DataRequired(), Length(max=11), Regexp(regex='^\d+$', message="Phone number must contain only digits")])
    email = StringField('Email', validators= [DataRequired(), Email()])
    password = PasswordField('Password', validators=[DataRequired(), Length(min=6)])
    Hash_password = PasswordField('Confirm Password', validators=[DataRequired(), EqualTo('password', message="Password mismatch")])
    submit = SubmitField('Register')


######################################################################
#database Model Design
#......................flask_model....................................
#farmrs model

class Farmer(UserMixin, db.Model):
    __tablename__ = 'farmers'

    # Primary key using UUID
    serial_number = db.Column(db.Integer, unique=True, nullable=False)
    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    # Basic Information
    email = db.Column(db.String(150), unique=True, nullable=False)
    password = db.Column(db.String(128), nullable=True)
    farmer_firstname = db.Column(db.String(150), nullable=False)
    farmer_lastname = db.Column(db.String(150), nullable=False)
    farmer_middlename = db.Column(db.String(150), nullable=True)
    farm_location = db.Column(db.String(200), nullable=False)
    date_of_birth = db.Column(db.Date, nullable=False)
    farmer_gender = db.Column(db.String(10))
    phone_number = db.Column(db.String(15), unique=True, nullable=False)
    bvn_number = db.Column(db.String(11), unique=True, nullable=False)
    id_card_type = db.Column(db.String(50), nullable=False)
    id_card_number = db.Column(db.String(50), unique=True, nullable=False)
    farm_size = db.Column(db.Float, nullable=False)
    gps_coordinates = db.Column(db.String(100), nullable=True)
    ward = db.Column(db.String(100), nullable=True)
    local_govt_council = db.Column(db.String(100), nullable=True)
    passport_photo = db.Column(db.String(255), nullable=True)
    role = db.Column(db.String(20), default='farmer')

    # Timestamps for auditing
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    # Foreign keys to track admin/app_admin actions with UUID references
    created_by_admin_id = db.Column(UUID, db.ForeignKey('admins.id'))
    updated_by_admin_id = db.Column(UUID, db.ForeignKey('admins.id'))
    deleted_by_admin_id = db.Column(UUID, db.ForeignKey('admins.id'))

    created_by_app_admin_id = db.Column(UUID, db.ForeignKey('app_admins.id'))
    updated_by_app_admin_id = db.Column(UUID, db.ForeignKey('app_admins.id'))
    deleted_by_app_admin_id = db.Column(UUID, db.ForeignKey('app_admins.id'))

    # Relationships to Admin and AppAdmin models
    created_by_admin = db.relationship('Admin', foreign_keys=[created_by_admin_id], back_populates='managed_created_farmers')
    updated_by_admin = db.relationship('Admin', foreign_keys=[updated_by_admin_id], back_populates='managed_updated_farmers')
    deleted_by_admin = db.relationship('Admin', foreign_keys=[deleted_by_admin_id], back_populates='managed_deleted_farmers')

    created_by_app_admin = db.relationship('AppAdmin', foreign_keys=[created_by_app_admin_id], back_populates='managed_created_farmers')
    updated_by_app_admin = db.relationship('AppAdmin', foreign_keys=[updated_by_app_admin_id], back_populates='managed_updated_farmers')
    deleted_by_app_admin = db.relationship('AppAdmin', foreign_keys=[deleted_by_app_admin_id], back_populates='managed_deleted_farmers')
     
    # UserMixin properties
   
    @property
    def is_active(self):
        return True

    @property
    def is_authenticated(self):
        return True

    @property
    def is_anonymous(self):
        return False
    
    def get_id(self):
        return f'farmer_{self.id}'

    
########################################

# Admin Model
class Admin(UserMixin, db.Model):
    __tablename__ = 'admins'
    
    # Primary key using UUID
    serial_number = db.Column(db.Integer, unique=True, nullable=False)
    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    
    # Basic Information
    email = db.Column(db.String(150), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    phone_number = db.Column(db.String(15), nullable=False)
    admin_name = db.Column(db.String(150), nullable=False)
    
    # Role field
    role = db.Column(db.String(20), default='admin')
    
    # Relationships with Farmer model using specific foreign keys
    managed_created_farmers = db.relationship(
        'Farmer',
        foreign_keys='Farmer.created_by_admin_id',
        back_populates='created_by_admin'
    )
    managed_updated_farmers = db.relationship(
        'Farmer',
        foreign_keys='Farmer.updated_by_admin_id',
        back_populates='updated_by_admin'
    )
    managed_deleted_farmers = db.relationship(
        'Farmer',
        foreign_keys='Farmer.deleted_by_admin_id',
        back_populates='deleted_by_admin'
    )

    def get_id(self):
        return f'admin_{self.id}'
       

class AppAdmin(UserMixin, db.Model):
    __tablename__ = 'app_admins'
    
    # Primary key using UUID
    serial_number = db.Column(db.Integer, nullable=False)
    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
    
    # Basic Information
    email = db.Column(db.String(150), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    admin_name = db.Column(db.String(150), nullable=False)
    
    # Role field
    role = db.Column(db.String(20), default='app_admin')
    
    # Relationships with Farmer model using specific foreign keys
    managed_created_farmers = db.relationship(
        'Farmer',
        foreign_keys='Farmer.created_by_app_admin_id',
        back_populates='created_by_app_admin'
    )
    managed_updated_farmers = db.relationship(
        'Farmer',
        foreign_keys='Farmer.updated_by_app_admin_id',
        back_populates='updated_by_app_admin'
    )
    managed_deleted_farmers = db.relationship(
        'Farmer',
        foreign_keys='Farmer.deleted_by_app_admin_id',
        back_populates='deleted_by_app_admin'
    )

    def get_id(self):
        return f'app_admin_{self.id}'


class ActionLog(db.Model):
    __tablename__ = 'action_logs'
    
    # Primary key
    serial_number = db.Column(db.Integer, unique=True, nullable=False)
    id = Column(UUID(as_uuid=True), primary_key=True, default=uuid.uuid4)
     # Action details
    action_type = db.Column(db.String(50), nullable=False)  # e.g., "create", "edit", "delete"
    entity_type = db.Column(db.String(50), nullable=False)  # e.g., "Farmer", "Admin"
    entity_id = db.Column(UUID, nullable=False)  # UUID of the affected entity
    
    # Timestamp
    timestamp = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    
   # Foreign keys for tracking who performed the action
    performed_by_admin_id = db.Column(UUID, db.ForeignKey('admins.id'), nullable=True)
    performed_by_app_admin_id = db.Column(UUID, db.ForeignKey('app_admins.id'), nullable=True)

    # Relationships to Admin and AppAdmin models
    performed_by_admin = db.relationship('Admin', foreign_keys=[performed_by_admin_id])
    performed_by_app_admin = db.relationship('AppAdmin', foreign_keys=[performed_by_app_admin_id])


#...........................................................................#
#...........................flask_views.................................#

#login function.............

@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        email_or_phone = form.email_or_phone.data
        password = form.password.data

        # Try user login based on email or phone number
        user = None
        if "@" in email_or_phone:
            user = Farmer.query.filter_by(email=email_or_phone).first() or \
                   Admin.query.filter_by(email=email_or_phone).first()  or \
                   AppAdmin.query.filter_by(email=email_or_phone).first()
        else:
            user = Farmer.query.filter_by(email=email_or_phone).first() or \
                   Admin.query.filter_by(email=email_or_phone).first()  or \
                   AppAdmin.query.filter_by(email=email_or_phone).first()
       
        # Check login credentials if user found
        if user and check_password_hash(user.password, password):
            login_user(user)
            session['user'] = user.id  # Store user ID in session
            session['role'] = 'app_admin' if isinstance(user, AppAdmin) else 'admin' if isinstance(user, Admin) else 'farmer'
            flash(f'Login successful as {session["role"]}!', 'success')

            # Redirect based on user role
            if isinstance(user, AppAdmin):
                # app.logger.warning('this app_admin user', session['_user_id'])
                return redirect(url_for('app_admin_dashboard'))
            elif isinstance(user, Admin):
                return redirect(url_for('admin_dashboard'))
            elif isinstance(user, Farmer):
                return redirect(url_for('farmer_dashboard', farmer_id=str(user.id)))
            else:
                flash('Invalid user role', 'danger')
                return render_template('login.html', form=form)

        else:
            flash('Invalid email, phone number, or password', 'danger')
            return render_template('login.html', form=form)

    return render_template('login.html', form=form)



# Ensure to configure your upload folder in app config
def allowed_file(filename):
    allowed_extensions = {'csv', 'xls', 'xlsx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

# from datetime import datetime
################################ ##########################

@app.route('/upload_farmers', methods=['POST'])
@csrf.exempt
@login_required
def upload_farmers_file():
    # if not isinstance(current_user, AppAdmin):
    #     abort(403)  # Forbidden for non-Admins and non-AppAdmins

    if 'file' not in request.files:
        flash('No file part', 'error')
        return redirect(url_for('app_admin_dashboard'))

    file = request.files['file']
    
    if file.filename == '':
        flash('No selected file', 'error')
        return redirect(url_for('app_admin_dashboard'))

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        success_count = 0
        failure_info = []
        file_duplicates = []
        existing_db_records = []

        try:
            # Read the file
            if filename.endswith('.csv'):
                df = pd.read_csv(filepath)
            elif filename.endswith('.xlsx'):
                df = pd.read_excel(filepath)
            else:
                flash('Unsupported file type', 'error')
                return redirect(url_for('app_admin_dashboard'))

            # Normalize column names
            df.columns = df.columns.str.strip().str.lower()
            
            # Define required columns
            required_columns = {
                'first name', 'last name', 'email', 'phone number', 'bvn number',
                'farm location', 'farm size(ha)', 'gender', 'gps coordinates',
                'id card no', 'type of identity (id) card', 'date of birth'
            }
            
            # Check if required columns are present
            missing_columns = required_columns - set(df.columns)
            if missing_columns:
                flash(f'Missing columns: {", ".join(missing_columns)}', 'error')
                return redirect(url_for('app_admin_dashboard'))
            
            # Prepare data for processing
            farmers_data = df.to_dict(orient='records')
            existing_emails = {farmer.email: farmer for farmer in Farmer.query.all()}
            existing_bvns = {farmer.bvn_number for farmer in Farmer.query.all()}
            existing_phone_numbers = {farmer.phone_number for farmer in Farmer.query.all()}

            # Process each record
            for data in farmers_data:
                email = str(data.get('email', '')).strip()
                bvn_number = str(data.get('bvn number', '')).strip()
                phone_number = str(data.get('phone number', '')).strip()

                # Check for duplicates in the database
                if email in existing_emails or bvn_number in existing_bvns or phone_number in existing_phone_numbers:
                    existing_db_records.append(data)
                    continue

                # Check for duplicates within the file
                if email in {row.get('email') for row in file_duplicates} or \
                   bvn_number in {row.get('bvn number') for row in file_duplicates} or \
                   phone_number in {row.get('phone number') for row in file_duplicates}:
                    file_duplicates.append(data)
                    continue

                # Handle missing required fields
                farmer_firstname = data.get('first name')
                farmer_lastname = data.get('last name')
                date_of_birth = data.get('date of birth')
                farm_location = data.get('farm location')

                if not all([farmer_firstname, farmer_lastname, date_of_birth, farm_location]):
                    failure_info.append({
                        'row': data,
                        'error': 'Required field missing'
                    })
                    continue
 
                try:
                    if isinstance(date_of_birth, str):
                        date_of_birth = parser.parse(date_of_birth).date()
                    elif isinstance(date_of_birth, datetime):
                        date_of_birth = date_of_birth.date()  # Ensure it's a date object
                    else:
                        raise ValueError('Invalid date format')
                except ValueError:
                    failure_info.append({
                        'row': data,
                        'error': 'Invalid date of birth format'
                    })
                    continue


                # Ensure the phone number is 11 digits and starts with a leading zero
                if len(phone_number) == 10:  # Likely missing leading zero
                    phone_number = '0' + phone_number
                elif len(phone_number) != 11 or not phone_number.startswith('0') or not phone_number.isdigit():
                    failure_info.append({
                        'row': data,
                        'error': 'Phone number must be 11 digits and start with a leading Zero'
                    })
                    continue

                # Validate BVN length and content
                if not bvn_number or len(bvn_number) != 11 or not bvn_number.isdigit():
                    failure_info.append({
                        'row': data,
                        'error': 'BVN number must be exactly 11 digits and numeric'
                    })
                    continue

                # Automatically set the password as the BVN, and hash it
                hashed_password = generate_password_hash(bvn_number)

                # Fetch the last serial number
                last_serial = db.session.query(func.max(Farmer.serial_number)).scalar() or 0

                # Create and add new farmer
                try:
                    new_farmer = Farmer(
                        email=email,
                        farmer_firstname=farmer_firstname,
                        farmer_lastname=farmer_lastname,
                        farm_location=farm_location,
                        date_of_birth=date_of_birth,
                        farmer_gender=data.get('gender'),
                        phone_number=phone_number,  # Ensure phone number is correct
                        bvn_number=bvn_number,
                        password=hashed_password,
                        id_card_type=data.get('type of identity (id) card'),
                        id_card_number=data.get('id card no'),
                        farm_size=data.get('farm size(ha)'),
                        gps_coordinates=data.get('gps coordinates'),
                        created_at=datetime.utcnow(),
                        updated_at=datetime.utcnow(),
                        created_by_admin_id=current_user.id,
                        # Assign the next available serial number
                        serial_number=last_serial + 1
                    )
                    db.session.add(new_farmer)
                    success_count += 1

                except IntegrityError:
                    db.session.rollback()
                    failure_info.append({
                        'row': data,
                        'error': 'Duplicate record or database error'
                    })
                except Exception as e:
                    db.session.rollback()
                    failure_info.append({
                        'row': data,
                        'error': f'Error creating farmer: {str(e)}'
                    })

            # Commit changes to the database
            db.session.commit()
        except Exception as e:
            flash(f'An unexpected error occurred: {str(e)}', 'error')
        finally:
            if filepath and os.path.exists(filepath):
                os.remove(filepath)

        # Flash results and redirect
        flash(f'{success_count} farmers added successfully.', 'success')
        if failure_info or file_duplicates or existing_db_records:
            session['failure_info'] = failure_info
            session['file_duplicates'] = file_duplicates
            session['existing_db_records'] = existing_db_records
            return redirect(url_for('farmer_upload_results'))

        return redirect(url_for('app_admin_dashboard'))

    flash('Invalid file type', 'error')
    return redirect(url_for('app_admin_dashboard'))


#############################################################################

######################################

@app.route('/farmer_upload_results')
@login_required
def farmer_upload_results():

    # if not isinstance(current_user, AppAdmin):
    #     abort(403)  # Forbidden for non-Admins and non-AppAdmins
 
    """
    Displays the results of the farmer file upload, including successes, duplicates, and errors.
    Also generates and serves a .docx file if requested.
    """
    success_count = session.get('success_count', 0)
    duplicates = session.get('file_duplicates', [])
    failure_info = session.get('failure_info', [])
    existing_db_records = session.get('existing_db_records', [])

    # If 'download' is in the query parameters, generate the .docx file
    if 'download' in request.args:
        doc = Document()
        doc.add_heading('Farmer Upload Errors', 0)

        if failure_info:
            doc.add_heading('Failures:', level=1)
            for failure in failure_info:
                doc.add_paragraph(f"Row: {failure['row']}")
                doc.add_paragraph(f"Error: {failure['error']}")
                doc.add_paragraph('')

        if duplicates:
            doc.add_heading('File Duplicates:', level=1)
            for duplicate in duplicates:
                doc.add_paragraph(str(duplicate))
                doc.add_paragraph('')

        if existing_db_records:
            doc.add_heading('Existing Database Records:', level=1)
            for existing in existing_db_records:
                doc.add_paragraph(str(existing))
                doc.add_paragraph('')

        # Save the document to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Send the document as a download
        return send_file(doc_io, as_attachment=True, download_name='farmer_upload_errors.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # Otherwise, render the results page
    return render_template('farmer_upload_results.html',
                           success_count=success_count,
                           duplicates=duplicates,
                           failure_info=failure_info)


#add a farmer
@app.route('/add_farmer', methods=['GET', 'POST'])
@login_required
def add_farmer():

    filepath = None  # Initialize filepath to None
    form = FarmerForm()
     
    #or isinstance(current_user, Farmer)
    # #Authorized users  
    # if not (isinstance(current_user, Admin) or isinstance(current_user, AppAdmin)):
    #     abort(403)

    if form.validate_on_submit():
        
        existing_entry = (Farmer.query.filter_by(bvn_number=form.bvn_number.data).first() or
                          Farmer.query.filter_by(id_card_number=form.id_card_number.data).first())
        if existing_entry:
            flash('An entry with this BVN or ID card number already exists.', 'danger')
            return redirect(url_for('add_farmer'))

         # Save file and get filename
        if form.passport_photo.data:
            file = form.passport_photo.data

            # Ensure the file object is valid and has a filename
            if isinstance(file, FileStorage) and file.filename:
                filename = secure_filename(file.filename)
                if filename:
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    file.save(filepath)
                    file.passport_photo = filename  # Save the filename in the farmer object
            # else:
            #     flash('Invalid file upload', 'danger')
                      # Ensure BVN is provided
            if not form.bvn_number.data:
                flash('BVN number is required to set the password.', 'danger')
                return redirect(url_for('add_farmer'))  # Adjust this to your appropriate redirect

             # Fetch the last serial number   
            last_serial = db.session.query(func.max(Farmer.serial_number)).scalar() or 0
            # Create a new farmer entry
            new_farmer = Farmer(
                email=form.email.data,
                farmer_firstname=form.farmer_firstname.data,
                farmer_lastname=form.farmer_lastname.data,
                farmer_middlename=form.farmer_middlename.data,
                date_of_birth=form.date_of_birth.data,
                # Automatically set the password as the BVN , and hash it
                password=generate_password_hash(form.bvn_number.data),
                phone_number=form.phone_number.data,
                bvn_number=form.bvn_number.data,
                id_card_type=form.id_card_type.data,
                id_card_number=form.id_card_number.data,
                farm_size=form.farm_size.data,
                gps_coordinates=form.gps_coordinates.data,
                farm_location=form.farm_location.data,
                ward=form.ward.data,
                local_govt_council=form.local_govt_council.data,
                created_by_admin_id=current_user.id if isinstance(current_user, Admin) else None,
                created_by_app_admin_id=current_user.id if isinstance(current_user, AppAdmin) else None,
                passport_photo=filename,  # Store filename instead of file field object
               
                # Assign the next available serial number
                serial_number=last_serial + 1
            )
            
            try:
                db.session.add(new_farmer)
                db.session.commit()
                
                # Optional: Move the file to a permanent location or perform other operations
                # os.rename(filepath, os.path.join(app.config['PERMANENT_UPLOAD_FOLDER'], filename))
                
                flash('Farmer added successfully!', 'success')
                return redirect(url_for('add_farmer'))
            except Exception as e:
                db.session.rollback()
                flash(f'Error adding farmer: {str(e)}', 'danger')
                return redirect(url_for('add_farmer'))
            # finally:
            #     # Remove the image after saving to path
            #     if filepath and os.path.exists(filepath):
            #         os.remove(filepath)

    return render_template('add_farmer.html', form=form)


#add_admin view function
@app.route('/add_admin', methods=['GET', 'POST'])
@login_required
def add_admin():

    form = AdminForm()

    # if not isinstance(current_user, AppAdmin):
    #     abort(403)  # Forbidden for non-AppAdmins

    if form.validate_on_submit():
        # Check if the admin already exists by email or phone number
        existing_admin = Admin.query.filter(
            (Admin.email == form.email.data) | 
            (Admin.phone_number == form.phone_number.data)
        ).first()

        if existing_admin:
            flash('Admin already exists.', 'danger')
        else:
            # Hash the password before storing it in the database
            hashed_password = generate_password_hash(form.password.data)

            # Fetch the last serial number
            last_serial = db.session.query(func.max(Admin.serial_number)).scalar() or 0


            # Create a new admin instance
            new_admin = Admin(
                email=form.email.data,
                admin_name=form.admin_name.data,
                phone_number=form.phone_number.data,
                password=hashed_password,  # Assign the hashed password
                

                # Assign the next available serial number
                serial_number=last_serial + 1
            )

            try:
                # Add the new admin to the database
                db.session.add(new_admin)
                db.session.commit()

                flash('Admin added successfully!', 'success')
                return redirect(url_for('add_admin'))

            except Exception as e:
                db.session.rollback()
                flash(f'Error adding Admin: {str(e)}', 'danger')
                return redirect(url_for('add_admin'))

    return render_template('add_admin.html', form=form)

 #editing a farmer
 ####.................barley lord of moriah...................#########

@app.route('/farmer/dashboard/<int:farmer_id>', methods=['GET', 'POST'])
@login_required
def farmer_dashboard(farmer_id):

     #---- Fetch farmer and perform access control ----
    farmer = Farmer.query.get_or_404(farmer_id)

    # Check access control based on user role and farmer ID
    if current_user.role == 'farmer':
        # Only allow the logged-in farmer to access their own profile
        if current_user.id != farmer.id:
            abort(403)  # Unauthorized access
    elif current_user.role in ('admin', 'app_admin'):
        # Admins and AppAdmins can access any farmer's dashboard
        pass
    else:
        abort(403)  # Unauthorized access for other roles

    # ---- Form handling for POST requests ----
    form = EditFarmerForm(obj=farmer)
    if form.validate_on_submit():
  
        # Check for duplicate entries by BVN or ID card
        existing_entry = (Farmer.query.filter(Farmer.id != farmer.id)
                          .filter((Farmer.bvn_number == form.bvn_number.data) | (Farmer.id_card_number == form.id_card_number.data))
                          .first())
        if existing_entry:
            flash('An entry with this BVN or ID card number already exists.', 'danger')
            return redirect(url_for('farmer_dashboard', farmer_id=farmer.id))  # Redirect to the same page to show errors
        
        # Save file and get filename
        if form.passport_photo.data:
            file = form.passport_photo.data

            # Ensure the file object is valid and has a filename
            if isinstance(file, FileStorage) and file.filename:
                filename = secure_filename(file.filename)
                if filename:
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    file.save(filepath)
                    farmer.passport_photo = filename  # Save the filename in the farmer object
            # else:
            #     flash('Invalid file upload', 'danger')

           # Update farmer details based on user type
        update_farmer_details(current_user, farmer, form)

        # Fetch the last serial number
        last_serial = db.session.query(func.max(Farmer.serial_number)).scalar() or 0

        try:
            db.session.commit()
            # ... log action ...
            action_log = ActionLog(
                action_type="edit",
                entity_type="Farmer",
                entity_id=farmer.id,
                performed_by_admin_id=current_user.id if isinstance(current_user, Admin) else None,
                performed_by_app_admin_id=current_user.id if isinstance(current_user, AppAdmin) else None,
                
                # Assign the next available serial number
                serial_number=last_serial + 1
            )

            db.session.add(action_log)
            db.session.commit()  # Commit the log

            flash('Farmer details updated successfully!', 'success')
            return redirect(url_for('farmer_dashboard', farmer_id=farmer.id))
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating farmer details: {str(e)}', 'danger')

    # ---- Prepare context for the template ----
    filepath = None  # Initialize filepath

    return render_template('farmer_dashboard.html', form=form, farmer=farmer, current_user=current_user)


def update_farmer_details(user, farmer, form):
    """
    Updates farmer details based on the current user's role.
    """
    # ... update other farmer attributes ...
    # Update farmer details
    farmer.email = form.email.data
    farmer.farmer_firstname = form.farmer_firstname.data
    farmer.farmer_lastname = form.farmer_lastname.data
    farmer.phone_number = form.phone_number.data
    farmer.farm_location = form.farm_location.data
    farmer.farm_size = form.farm_size.data
    farmer.gps_coordinates = form.gps_coordinates.data
    farmer.id_card_type = form.id_card_type.data
    farmer.bvn_number = form.bvn_number.data
    farmer.id_card_number = form.id_card_number.data

    # Set updated_by fields based on user type
    if isinstance(user, Admin):
        farmer.updated_by_admin_id = user.id
    elif isinstance(user, AppAdmin):
        farmer.updated_by_app_admin_id = user.id
    else:  # Farmer editing their own profile
        farmer.updated_by_farmer_id = user.id


###################....................................#####################

@app.route('/admin/<int:admin_id>')

def view_admin(admin_id):
    
    admin = Admin.query.get_or_404(admin_id)
    #  users = User.query.all()
    return render_template('view_admin.html', admin=admin)


########################################################################
#............edit admin............#######
@app.route('/edit_admin/<int:admin_id>', methods=['GET', 'POST'])
@login_required
def edit_admin(admin_id):

    form = AdminForm()
    
    admin = Admin.query.get_or_404(admin_id)

    # if not isinstance(current_user, AppAdmin):
        # abort(403)  # Only Admins and AppAdmins can edit farmers

    if request.method == 'POST':
        admin.admin_name = request.form['name']
        admin.email = request.form['email']
        admin.phone_number = request.form['phone']
        admin.password = generate_password_hash(request.form['password'])
        admin.updated_at = datetime.utcnow()
        
        try:
            db.session.commit()
            flash('Admin updated successfully!', 'success')
            return redirect(url_for('app_admin_dashboard'))
        
        except Exception as e:

            flash('Error updating admin.', 'danger')
            return redirect(url_for('edit_admin', admin_id=admin.id))
    
    return render_template('edit_admin.html', admin=admin, form=form)


#####################################################################

@app.route('/delete_admin/<int:admin_id>', methods=['POST'])
@login_required
def delete_admin(admin_id):
    admin = Admin.query.get_or_404(admin_id)

    # if not isinstance(current_user, AppAdmin):
    #     abort(403)  # Forbidden for non-Admins and non-AppAdmins

    # # Only allow AppAdmin users to perform this action
    # if current_user.role != 'app_admin':
    #     abort(403)  # Access forbidden if not app_admin
    
    # Fetch the last serial number
    last_serial = db.session.query(func.max(ActionLog.serial_number)).scalar() or 0
    try:
        db.session.delete(admin)
        db.session.commit()
        # Log the action
        action_log = ActionLog(
            action_type="delete",
            entity_type="Admin",
            entity_id=admin.id,
            # performed_by_admin_id=current_user.id if isinstance(current_user, Admin) else None,
            # performed_by_app_admin_id=current_user.id if isinstance(current_user, AppAdmin) else None
           
           
           # Assign the next available serial number 
            serial_number=last_serial + 1
        )
        db.session.add(action_log)
        db.session.commit()

        # Flash message for successful deletion
        flash('Admin deleted successfully.', 'success')
    except Exception as e:
        # Rollback the session if there is an error
        db.session.rollback()
        flash('An error occurred while deleting the Admin.', 'danger')

    return '', 204  # No content response


#delete farmer
@app.route('/delete_farmer/<int:farmer_id>', methods=['POST'])
@login_required
def delete_farmer(farmer_id):
    farmer = Farmer.query.get_or_404(farmer_id)

    # if not isinstance(current_user, Admin) and not isinstance(current_user, AppAdmin):
    #     abort(403)  # Only Admins and AppAdmins can delete farmers

    # Fetch the last serial number
    last_serial = db.session.query(func.max(Farmer.serial_number)).scalar() or 0

    # Attempt to delete the farmer
    try:
        db.session.delete(farmer)
        db.session.commit()


        # Log the action
        action_log = ActionLog(
            action_type="delete",
            entity_type="Farmer",
            entity_id=farmer.id,
            performed_by_admin_id=current_user.id if isinstance(current_user, Admin) else None,
            performed_by_app_admin_id=current_user.id if isinstance(current_user, AppAdmin) else None,

            # Assign the next available serial number
            serial_number=last_serial + 1
        )
        db.session.add(action_log)
        db.session.commit()

        # Flash message for successful deletion
        flash('Farmer deleted successfully.', 'success')
    except Exception as e:
        # Rollback the session if there is an error
        db.session.rollback()
        flash('An error occurred while deleting the farmer.', 'danger')

    return '', 204  # No content response

############################################################################

@app.route('/search_admin', methods=['GET', 'POST'])
@csrf.exempt
# @login_required
def search_admin():

    if not isinstance(current_user, AppAdmin):
        abort(403)  # Forbidden for non-AppAdmins

    # Initialize an empty list to hold search results
    admin_results = []

    if request.method == 'POST':
        admin_search = request.form.get('admin_search').strip()

        if admin_search:
            admin_results = Admin.query.filter(
                (Admin.email.ilike(f'%{admin_search}%')) | 
                (Admin.admin_name.ilike(f'%{admin_search}%'))
            ).all()

            if not admin_results:
                flash('Admin not found!!', 'danger')
        else:
            flash('Please enter a search term', 'warning')


    return render_template(
        'app_admin_dashboard.html',
        admin_results=admin_results)


@app.route('/search_farmer', methods=['GET', 'POST'])
# @login_required
@csrf.exempt
def search_farmer():
    if not isinstance(current_user, Admin) and not isinstance(current_user, AppAdmin):
        abort(403)  # Forbidden for non-Admins and non-AppAdmins
 
    # Initialize an empty list to hold search results
    farmer_results = []

    if request.method == 'POST':
        farmer_search = request.form.get('farmer_search').strip()

        if farmer_search:
            farmer_results = Farmer.query.filter(
                (Farmer.phone_number.ilike(f'%{farmer_search}%')) | 
                (Farmer.bvn_number.ilike(f'%{farmer_search}%'))|
                (Farmer.id_card_number.ilike(f'%{farmer_search}%'))
            ).all()

            if not farmer_results:
                flash('No results found!!!', 'danger')
        else:
            flash('Please enter a search term!!', 'warning')

            
    # return render_template('app_admin_dashboard.html',
    #                        farmer_results=farmer_results)
            
        # Render the dashboard template
        if isinstance(current_user, AppAdmin):
            return render_template(
                'app_admin_dashboard.html',
                farmer_results=farmer_results, 
        )
        if isinstance(current_user, Admin):
            return render_template(
                'admin_dashboard.html',
                farmer_results=farmer_results, 
        )
        else:
            return redirect(url_for('login'))
        # isinstance(current_user, Admin)
    return redirect(request.url)


#dashboard for all the 3 users.....................#####
@app.route('/app_admin/dashboard')
@app_admin_required
def app_admin_dashboard():

    app.logger.warning('Current user: %s', current_user)
    app.logger.warning("Session user data:", session.get("user"))
   
    # Initialize forms
    admin_form = SearchAdminForm()
    farmer_form = SearchFarmerForm()

    # Initialize empty results lists
    admin_results = []
    farmer_results = []

    # Handle admin search form submission
    if admin_form.validate_on_submit():
        admin_search = admin_form.admin_search.data
        admin_results = Admin.query.filter(
            (Admin.email.ilike(f'%{admin_search}%')) | 
            (Admin.admin_name.ilike(f'%{admin_search}%'))
        ).all()

        if not admin_results:
            flash('No results found for the admin search query', 'danger')

    # Handle farmer search form submission
    if farmer_form.validate_on_submit():
        farmer_search = farmer_form.farmer_search.data
        farmer_results = Farmer.query.filter(
            (Farmer.phone_number.ilike(f'%{farmer_search}%')) | 
            (Farmer.bvn_number.ilike(f'%{farmer_search}%'))
        ).all()

        if not farmer_results:
            flash('No results found for the farmer search query', 'danger')

    # Render the dashboard template based on the user type
    if isinstance(current_user, AppAdmin):
        return render_template(
            'app_admin_dashboard.html',
            farmer_results=farmer_results, 
            admin_results=admin_results, 
            farmer_form=farmer_form, 
            admin_form=admin_form
            )
    
    elif isinstance(current_user, Admin):
        return render_template(
            'admin_dashboard.html',
            farmer_results=farmer_results, 
            admin_results=admin_results, 
            farmer_form=farmer_form, 
            admin_form=admin_form
        )
    
    else:
        return redirect(url_for('login'))  # Redirect to login if user is neither AppAdmin nor Admin


@app.route('/admin/dashboard')
# @login_required
def admin_dashboard():

    if not isinstance(current_user, Admin) | isinstance(current_user, AppAdmin):
        abort(403)  # Forbidden

    return render_template('admin_dashboard.html')



######........................................................###
#########..........view all farmers and admin.............#######

@app.route('/preview_farmers/<int:farmer_id>', methods=['GET'])
# @login_required
def view_farmers(farmer_id=None):
    form = FarmerForm()

    # If a specific farmer_id is provided, fetch the corresponding farmer
    if farmer_id:
        farmer = Farmer.query.get(farmer_id)
        if not farmer:
            abort(404)  # Farmer not found
        return render_template('preview_farmer.html', form=form, farmer=farmer)

    # If no farmer_id is provided, and the current user is a farmer, show their profile
    if isinstance(current_user, Farmer):
        farmer = Farmer.query.filter_by(id=current_user.id).first()
        if farmer:
            return render_template('preview_farmer.html', form=form, farmer=farmer)
        else:
            abort(404)  # Current farmer not found

    # If the current user is an Admin or AppAdmin, display a list of all farmers
    if isinstance(current_user, Admin) or isinstance(current_user, AppAdmin):
        farmers = Farmer.query.all()
        return render_template('view_all_farmers.html', form=form, farmers=farmers)

    return redirect(url_for('login'))  # If not authenticated



##################################################
#logout
@app.route('/logout')
def logout():
    logout_user()
    app.logger.warning(session.keys())
    app.logger.warning(session.get('_user_id', 'no user id'))
    return redirect(url_for('login'))

   
if __name__ == '__main__':
    with app.app_context():
        
        db.create_all()
        #create default admin
        create_default_admin()
    app.run(host='0.0.0.0', debug=True)
